"""Fachliche End-to-End-Fälle gegen ECHTE Umsysteme (STT + LLM).

Nur auf Promotion (Marker `promotion`), nicht bei jedem Build – Testkonzept §9.
Fliesst eine echte Diktat-Aufnahme durch die ganze Kette (Audio → STT → Interview
→ LLM-Extraktion → PIA) und prüft INVARIANTEN (§10), nicht exakte Gleichheit.

Aktiviert sich automatisch, sobald
  1. eine Audio-Aufnahme unter tests/e2e/fixtures/<fall-id>_*.<ext> liegt und
  2. STT_API_KEY und ANTHROPIC_API_KEY gesetzt sind.
Sonst wird der Fall übersprungen (kein Fehlschlag) – siehe tests/e2e/README.md.
"""
import glob
import json
import os
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document

from app.config import Config
from app.shared.database import SessionLocal
from tests.e2e.invarianten import (
    erfundene_fundstelle,
    hermes_term_violations,
    projekttyp_ist_gueltig,
)

_MIME = {".webm": "audio/webm", ".wav": "audio/wav", ".mp3": "audio/mpeg",
         ".m4a": "audio/mp4", ".ogg": "audio/ogg", ".flac": "audio/flac"}


def _fixture_dirs():
    """Suchreihenfolge für Audio-Fixtures. E2E_FIXTURES_DIR zeigt auf einen
    GESCHÜTZTEN Ort auf dem Build-Agent (Aufnahmen liegen NIE im öffentlichen
    Repo); lokal dient der gitignorierte Repo-Ordner als Fallback."""
    dirs = []
    env = os.environ.get("E2E_FIXTURES_DIR")
    if env:
        dirs.append(Path(env))
    dirs.append(Path(__file__).parent / "fixtures")
    return dirs


def _audio_fixture(fall_id):
    for d in _fixture_dirs():
        treffer = sorted(glob.glob(str(d / f"{fall_id}_*")))
        treffer = [t for t in treffer if Path(t).suffix.lower() in _MIME]
        if treffer:
            return Path(treffer[0])
    return None


def _app_mit_echten_diensten():
    """App mit In-Memory-DB, aber echten API-Keys aus der Umgebung."""
    class _Cfg(Config):
        DATABASE_URL = "sqlite:///:memory:"
        SECRET_KEY = "e2e"
    SessionLocal.remove()
    from app.factory import create_app
    app = create_app(_Cfg)
    SessionLocal.remove()
    return app


# --------------------------------------------------------------------------- #
# Selbsttest des Orakels – läuft immer (verifiziert die Invarianten-Logik)     #
# --------------------------------------------------------------------------- #

def test_orakel_erkennt_hermes_verstoss():
    assert hermes_term_violations("Der Projektauftrag wurde erteilt.")
    assert hermes_term_violations("Der Steuerungsausschuss tagt.")
    assert hermes_term_violations("Ein Phasenbericht liegt vor.")
    # HERMES-konform: keine Verstösse
    assert hermes_term_violations("Der Durchführungsauftrag und der Projektausschuss.") == []


def test_orakel_erkennt_erfundene_fundstelle():
    assert erfundene_fundstelle("gemäss NG 236.1 ist ...")
    assert erfundene_fundstelle("nach SR 172.010 gilt ...")
    assert not erfundene_fundstelle("gemäss der geltenden Verordnung")


def test_orakel_projekttyp_gueltigkeit():
    assert projekttyp_ist_gueltig(None)
    assert projekttyp_ist_gueltig("e_government_portal")
    assert not projekttyp_ist_gueltig("irgendwas_erfundenes")


# --------------------------------------------------------------------------- #
# Der eigentliche E2E-Fall – nur auf Promotion, gegen echte Dienste            #
# --------------------------------------------------------------------------- #

def _testfall_ids():
    """IDs der fachlichen Testfälle aus dem Katalog (YAML). So laufen auf Promotion
    automatisch ALLE Fälle, für die eine Aufnahme hinterlegt ist – nicht nur einer."""
    pfad = Path(__file__).resolve().parents[1] / "fachlich" / "hermes_pia_testfaelle.yaml"
    try:
        import yaml
        data = yaml.safe_load(pfad.read_text(encoding="utf-8")) or {}
    except Exception:
        return ["pia-fachlich-0003"]
    faelle = data.get("testfaelle") or []
    ids = [f.get("id") for f in faelle if isinstance(f, dict) and f.get("id")]
    return ids or ["pia-fachlich-0003"]


@pytest.mark.promotion
@pytest.mark.parametrize("fall_id", _testfall_ids())
def test_e2e_fachfall_diktat_bis_pia(fall_id):
    """Jeder fachliche Fall MIT hinterlegter Aufnahme fliesst als echter Durchlauf
    (Aufnahme → STT → LLM → PIA) und wird gegen die Invarianten (§10) geprüft.
    Fälle ohne Aufnahme werden übersprungen (kein Fehlschlag)."""
    audio = _audio_fixture(fall_id)
    if audio is None:
        pytest.skip(f"Keine Audio-Aufnahme unter tests/e2e/fixtures/{fall_id}_* – siehe README.")
    if not os.environ.get("ANTHROPIC_API_KEY"):
        pytest.skip("ANTHROPIC_API_KEY fehlt – echter LLM-Lauf nicht möglich.")

    app = _app_mit_echten_diensten()
    isvc, gen = app.interview_service, app.generation_service
    if not getattr(app.transcriber, "available", False):
        pytest.skip("STT nicht konfiguriert (STT_API_KEY) – echter Transkriptionslauf nicht möglich.")

    with app.app_context():
        # 1) Echte Transkription der Aufnahme
        data = audio.read_bytes()
        transkript = app.transcriber.transcribe(
            data, filename=audio.name, mimetype=_MIME[audio.suffix.lower()])
        assert transkript and transkript.strip(), "STT lieferte keinen Text."

        # 2) Interview: Ausgangslage aus dem Diktat, echte LLM-Verarbeitung
        org = app.auth_service.create_org("E2E-Org")
        session = isvc.start_session(
            method_id="hermes_pia", project_name="E2E Testprojekt",
            org_id=org.id, created_by="Testperson", auftraggeber="Auftraggeberin")
        sid = session.id
        isvc.submit_answer(sid, transkript)

        session = isvc.get_session(sid)
        answers = json.loads(session.answers_json or "{}")

        # 3) Invarianten (Testorakel §10)
        assert projekttyp_ist_gueltig(session.project_type_id), \
            f"Ungültiger/geratener Projekttyp: {session.project_type_id!r}"

        ausg = (answers.get("ausgangslage") or {}).get("extracted") or {}
        ausg_text = ausg.get("text") if isinstance(ausg, dict) else ""
        assert ausg_text and ausg_text.strip(), "Ausgangslage wurde nicht erfasst."
        assert not erfundene_fundstelle(ausg_text), "Erfundene Gesetzes-Fundstelle in der Ausgangslage."

        # Nach der Ausgangslage folgt die Komplexitäts-/Folgefrage-Phase.
        state = isvc.current_state(session)
        assert state.get("phase") in ("followup", "question", "complete")

        # 4) PIA erzeugen und dokumentweit auf HERMES-Konformität prüfen
        metadata = {"projektname": "E2E Testprojekt", "projektleiter": "Testperson",
                    "auftraggeber": "Auftraggeberin", "version": "0.1"}
        buf = gen.generate("hermes_pia", answers, metadata)
        doc = Document(BytesIO(buf.getvalue() if isinstance(buf, BytesIO) else buf))
        volltext = "\n".join(p.text for p in doc.paragraphs)
        for t in doc.tables:
            for row in t.rows:
                for c in row.cells:
                    volltext += "\n" + c.text

        verstoesse = hermes_term_violations(volltext)
        assert not verstoesse, f"HERMES-Begriffsverstoss im PIA: {verstoesse}"
        assert not erfundene_fundstelle(volltext), "Erfundene Fundstelle im erzeugten PIA."
        assert any("Ausgangslage" in p.text for p in doc.paragraphs), \
            "Kapitel 'Ausgangslage' fehlt im erzeugten Dokument."

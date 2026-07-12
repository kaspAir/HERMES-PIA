"""Inkrement 3: Das Interview folgt der Kapitelstruktur einer hochgeladenen
Word-Vorlage (opt-in, 'live'). Ohne Vorlage bleibt alles wie mit der kanonischen
HERMES-Methode – das prüft implizit die gesamte übrige Suite mit.
"""
from io import BytesIO

import pytest
from docx import Document

from app.config import Config
from app.factory import create_app


@pytest.fixture
def app(tmp_path):
    from app.shared.database import SessionLocal

    db_path = str(tmp_path / "iv.db").replace("\\", "/")

    class _Cfg(Config):
        DATABASE_URL = "sqlite:///" + db_path
        SUPERADMIN_EMAIL = "betreiber@test.ch"
        SUPERADMIN_PASSWORD = "pw-super"
        SECRET_KEY = "test-secret"

    SessionLocal.remove()
    application = create_app(_Cfg)
    SessionLocal.remove()
    yield application
    SessionLocal.remove()


def _docx(headings):
    doc = Document()
    for level, text in headings:
        doc.add_heading(text, level=level)
        doc.add_paragraph("Text.")
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _start_session(app):
    auth = app.auth_service
    org = auth.create_org("Org")
    org_id = org.id  # sofort als int festhalten – Detached-Instance vermeiden
    auth.create_user("a@a.ch", "pw", role="org_admin", org_id=org_id,
                     can_read=True, can_write=True, can_delete=True)
    c = app.test_client()
    c.post("/login", data={"email": "a@a.ch", "password": "pw"})
    loc = c.post("/interview/start",
                 data={"project_name": "P", "projektleiter": "X"}).headers["Location"]
    sid = int(loc.rstrip("/").split("/")[-1])
    return org_id, sid


def test_ohne_vorlage_kanonische_struktur(app):
    _, sid = _start_session(app)
    isvc = app.interview_service
    ids = [s["id"] for s in isvc.section_summary(isvc.get_session(sid))]
    assert "ausgangslage" in ids
    assert "personalaufwand" in ids
    assert not any(i.startswith("custom_") for i in ids)


def test_interview_folgt_vorlage_live(app):
    org_id, sid = _start_session(app)
    isvc = app.interview_service
    session = isvc.get_session(sid)

    projekt = app.projekt_service.projekt_for_ergebnis(session.ergebnis_id)
    assert projekt is not None

    # HERMES-Variante: umbenannt + Zusatzkapitel, nur 3 Kapitel
    data = _docx([(1, "Ausgangssituation"), (1, "Zielsetzung"),
                  (1, "Datenschutzkonzept")])
    app.projekt_service.add_methoden_vorlage(
        "kunde.docx", data, org_id=org_id, projekt_id=projekt.id)

    # 'live': dieselbe Session folgt jetzt exakt der Vorlagenstruktur
    ids = [s["id"] for s in isvc.section_summary(isvc.get_session(sid))]
    assert ids == ["ausgangslage", "ziele", "custom_datenschutzkonzept"]

    # Die erste offene Frage kommt aus dem ersten Vorlagenkapitel (Ausgangslage)
    state = isvc.current_state(isvc.get_session(sid))
    assert state["phase"] == "question"
    assert state["section"]["id"] == "ausgangslage"


def test_projekt_vorlage_uebersteuert_org_im_interview(app):
    org_id, sid = _start_session(app)
    isvc = app.interview_service
    projekt = app.projekt_service.projekt_for_ergebnis(
        isvc.get_session(sid).ergebnis_id)

    # Org-Vorlage: nur Ausgangslage; Projekt-Vorlage: nur Ziele -> Projekt gewinnt
    app.projekt_service.add_methoden_vorlage(
        "org.docx", _docx([(1, "Ausgangslage")]), org_id=org_id, projekt_id=None)
    app.projekt_service.add_methoden_vorlage(
        "projekt.docx", _docx([(1, "Ziele")]), org_id=org_id, projekt_id=projekt.id)

    ids = [s["id"] for s in isvc.section_summary(isvc.get_session(sid))]
    assert ids == ["ziele"]


def test_generisches_kapitel_erfragen_und_speichern(app):
    org_id, sid = _start_session(app)
    isvc = app.interview_service
    projekt = app.projekt_service.projekt_for_ergebnis(
        isvc.get_session(sid).ergebnis_id)
    app.projekt_service.add_methoden_vorlage(
        "k.docx", _docx([(1, "Datenschutzkonzept")]),
        org_id=org_id, projekt_id=projekt.id)

    state = isvc.current_state(isvc.get_session(sid))
    assert state["section"]["id"] == "custom_datenschutzkonzept"
    # Das generische Kapitel ist Fliesstext mit Fragen
    assert state["section"]["type"] == "free_text"
    assert state["section"]["interview"]["questions"]

"""Kern-Engine Kundenvorlage: Kapitel auslesen + auf HERMES-Abschnitte mappen.

Prüft das Risikostück von Stufe 2 isoliert: erkennt die Engine die Kapitel einer
HERMES-Varianten-Vorlage (umbenannt, umsortiert, mit Zusatzkapiteln) verlässlich?
"""
from io import BytesIO

from docx import Document

from app.config import Config
from app.domains.method.template_structure import (
    ZIEL_GENERISCH,
    ZIEL_UNVERAENDERT,
    build_derived_method,
    build_method_from_mapping,
    extract_headings,
    match_canonical,
    propose_mapping,
)
from app.shared.config_loader import load_method


def _canonical():
    return load_method(Config.METHODS_DIR, "hermes_pia")


def _docx(headings):
    """Baut eine .docx mit den gegebenen (level, text)-Überschriften + etwas Text."""
    doc = Document()
    for level, text in headings:
        doc.add_heading(text, level=level)
        doc.add_paragraph("Platzhaltertext.")
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_extract_headings_liest_reihenfolge():
    data = _docx([(1, "Ausgangslage"), (1, "Ziele"), (2, "Unterpunkt")])
    heads = extract_headings(data)
    assert [h["text"] for h in heads] == ["Ausgangslage", "Ziele", "Unterpunkt"]
    assert heads[0]["level"] == 1 and heads[2]["level"] == 2


def test_extract_headings_robust_bei_muell():
    assert extract_headings(b"kein zip") == []
    assert extract_headings(b"") == []


def test_match_canonical_exakt_und_umlaut():
    sects = _canonical()["sections"]
    assert match_canonical("Ausgangslage", sects) == "ausgangslage"
    # führende Kapitelnummer wird ignoriert
    assert match_canonical("1 Ausgangslage", sects) == "ausgangslage"
    assert match_canonical("2.1 Ziele der Phase Initialisierung", sects) == "ziele"


def test_match_canonical_synonyme():
    sects = _canonical()["sections"]
    assert match_canonical("Ausgangssituation", sects) == "ausgangslage"
    assert match_canonical("Zielsetzung", sects) == "ziele"
    assert match_canonical("Budget", sects) == "kosten"
    assert match_canonical("Risikoanalyse", sects) == "risiken"


def test_match_canonical_unbekannt_gibt_none():
    sects = _canonical()["sections"]
    assert match_canonical("Datenschutzkonzept", sects) is None
    assert match_canonical("", sects) is None


def test_build_derived_erkennt_umbenannte_und_umsortierte_kapitel():
    # HERMES-Variante: umbenannt (Ausgangssituation/Zielsetzung/Budget),
    # umsortiert (Risiken vor Kosten) und ein Zusatzkapitel.
    data = _docx([
        (1, "Ausgangssituation"),
        (1, "Zielsetzung"),
        (1, "Datenschutzkonzept"),   # Zusatzkapitel -> generisch
        (1, "Risikoanalyse"),
        (1, "Budget"),
    ])
    method, report = build_derived_method(data, _canonical())

    reihenfolge = [s["id"] for s in method["sections"]]
    # Struktur der Vorlage bestimmt die Reihenfolge (Risiken vor Kosten!)
    assert reihenfolge == [
        "ausgangslage", "ziele", "custom_datenschutzkonzept", "risiken", "kosten",
    ]
    assert [m["section_id"] for m in report["matched"]] == \
        ["ausgangslage", "ziele", "risiken", "kosten"]
    assert report["generic"][0]["section_id"] == "custom_datenschutzkonzept"

    # Erkannte Kapitel behalten die volle kanonische Definition (Intelligenz!)
    kosten = next(s for s in method["sections"] if s["id"] == "kosten")
    assert kosten["type"] == "table"
    # Generisches Kapitel ist Fliesstext mit abgeleiteten Fragen
    ds = next(s for s in method["sections"] if s["id"] == "custom_datenschutzkonzept")
    assert ds["type"] == "free_text" and ds["generic"] is True
    assert len(ds["interview"]["questions"]) >= 1


def test_build_derived_meldet_fehlende_kanonische_kapitel():
    data = _docx([(1, "Ausgangslage"), (1, "Ziele")])
    _, report = build_derived_method(data, _canonical())
    # In der Vorlage fehlende Pflicht-/Inhaltskapitel werden gemeldet
    assert "risiken" in report["missing_canonical"]
    assert "personalaufwand" in report["missing_canonical"]
    assert "ausgangslage" not in report["missing_canonical"]


def test_container_und_strukturueberschriften_werden_uebersprungen():
    # Nachbildung der echten HERMES-Struktur: Titel, Inhaltsverzeichnis,
    # Container-L1 ("Einleitung", "Ressourcenbedarf") mit L2-Inhaltskapiteln,
    # Protokoll am Schluss.
    data = _docx([
        (1, "Projektinitialisierungsauftrag"),   # Dokumenttitel
        (1, "Inhaltsverzeichnis"),               # strukturell
        (1, "Einleitung"),                       # Container (nächstes L2)
        (2, "Referenzierte Dokumente"),
        (1, "Ausgangslage"),
        (1, "Ressourcenbedarf"),                 # Container
        (2, "Personalaufwand"),
        (2, "Kosten (in CHF inkl. MwSt.)"),
        (1, "Risiken"),
        (1, "Dokument-Protokoll"),               # strukturell
    ])
    method, report = build_derived_method(data, _canonical())
    ids = [s["id"] for s in method["sections"]]

    # Ausgangslage steht im Interview immer zuerst; die übrigen Kapitel behalten
    # die Vorlagenreihenfolge.
    assert ids == [
        "ausgangslage", "referenzierte_dokumente",
        "personalaufwand", "kosten", "risiken",
    ]
    # Weder Titel/Verzeichnis/Protokoll noch Container werden zu Kapiteln
    assert report["generic"] == []
    uebersprungen = {s["heading"] for s in report["skipped"]}
    assert {"Projektinitialisierungsauftrag", "Inhaltsverzeichnis",
            "Einleitung", "Ressourcenbedarf", "Dokument-Protokoll"} <= uebersprungen


def test_vorspann_hilfstexte_werden_nicht_erfragt():
    # KTZH-artig: fixe Hilfstexte VOR der Ausgangslage dürfen nicht zu Fragen werden.
    data = _docx([
        (1, "Änderungsverzeichnis"),
        (1, "Hinweise zum HERMES-Dokument"),
        (1, "Beschreibung"),
        (1, "Hinweise zur Anwendung des Dokumentes"),
        (1, "Ausgangslage"),
        (1, "Ziele"),
        (1, "Datenschutzkonzept"),   # NACH Ausgangslage -> echtes Zusatzkapitel
    ])
    method, report = build_derived_method(data, _canonical())
    ids = [s["id"] for s in method["sections"]]

    assert ids == ["ausgangslage", "ziele", "custom_datenschutzkonzept"]
    # Die Hilfstexte sind als Vorspann übersprungen, nicht generisch erfragt
    assert report["generic"] == [{"heading": "Datenschutzkonzept",
                                  "section_id": "custom_datenschutzkonzept"}]
    vorspann = {s["heading"] for s in report["skipped"] if s["grund"] == "vorspann"}
    assert {"Hinweise zum HERMES-Dokument", "Beschreibung",
            "Hinweise zur Anwendung des Dokumentes"} <= vorspann


def test_umbenanntes_inhaltskapitel_wird_nicht_verschluckt():
    # Ein echtes Kapitel, das wir nicht erkennen (umbenannt) und das VOR dem
    # ersten erkannten HERMES-Kapitel steht, darf NIE als Vorspann verworfen
    # werden – lieber als generisches Kapitel erfragen.
    data = _docx([(1, "Projektkontext"), (1, "Ziele"), (1, "Risiken")])
    method, report = build_derived_method(data, _canonical())
    ids = [s["id"] for s in method["sections"]]
    assert "custom_projektkontext" in ids
    assert not any(s["heading"] == "Projektkontext"
                   for s in report["skipped"] if s["grund"] == "vorspann")


def test_ausgangslage_wird_zuerst_erfragt():
    # Auch wenn die Vorlage Ausgangslage NICHT zuerst führt -> im Interview zuerst.
    data = _docx([(1, "Ziele"), (1, "Risiken"), (1, "Ausgangslage")])
    method, _ = build_derived_method(data, _canonical())
    assert method["sections"][0]["id"] == "ausgangslage"
    assert [s["id"] for s in method["sections"]] == ["ausgangslage", "ziele", "risiken"]


def test_build_derived_nutzt_question_gen_mit_fallback():
    data = _docx([(1, "Spezialkapitel")])

    # LLM-Ersatz liefert eigene Fragen
    method, _ = build_derived_method(
        data, _canonical(), question_gen=lambda t: [f"Frage zu {t}?"])
    sec = method["sections"][0]
    assert sec["interview"]["questions"] == ["Frage zu Spezialkapitel?"]

    # Fällt der Generator aus, greift der deterministische Fallback
    method2, _ = build_derived_method(
        data, _canonical(), question_gen=lambda t: (_ for _ in ()).throw(RuntimeError()))
    assert len(method2["sections"][0]["interview"]["questions"]) >= 1


def test_propose_mapping_seedet_vorschlag():
    data = _docx([
        (1, "Hinweise zur Anwendung des Dokumentes"),
        (1, "Ausgangssituation"),
        (1, "Zielsetzung"),
        (1, "Datenschutzkonzept"),
    ])
    vorschlag, missing = propose_mapping(data, _canonical())
    ziel = {v["heading"]: v["ziel"] for v in vorschlag}
    assert ziel["Hinweise zur Anwendung des Dokumentes"] == ZIEL_UNVERAENDERT
    assert ziel["Ausgangssituation"] == "ausgangslage"
    assert ziel["Zielsetzung"] == "ziele"
    assert ziel["Datenschutzkonzept"] == ZIEL_GENERISCH
    assert "risiken" in missing


def test_build_method_from_mapping_respektiert_bestaetigung():
    mapping = [
        {"heading": "Kontextkapitel", "ziel": "ausgangslage"},
        {"heading": "Unser Glossar", "ziel": "definitionen"},
        {"heading": "Spezielles", "ziel": ZIEL_GENERISCH},
        {"heading": "Fixe Einleitung", "ziel": ZIEL_UNVERAENDERT},
    ]
    method = build_method_from_mapping(_canonical(), mapping)
    ids = [s["id"] for s in method["sections"]]
    # Ausgangslage zuerst; unveraendertes Kapitel fehlt; generisches als custom_
    assert ids[0] == "ausgangslage"
    assert "definitionen" in ids
    assert "custom_spezielles" in ids
    assert not any("einleitung" in i for i in ids)
    # Die kanonische Section traegt die Original-Ueberschrift der Kundenvorlage
    ausg = next(s for s in method["sections"] if s["id"] == "ausgangslage")
    assert ausg["template_heading"] == "Kontextkapitel"

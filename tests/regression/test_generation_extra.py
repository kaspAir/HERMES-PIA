"""Beweist: Enddatum gefüllt, 'Projektauftrag' korrigiert, Kap.-6-Spalten passen."""
from docx import Document
from lxml import etree

from app.config import get_config
from app.domains.generation.service import GenerationService, W, _max_termin
from app.domains.method.service import MethodService


def _gen():
    cfg = get_config()
    return GenerationService(MethodService(cfg.METHODS_DIR))


def _full_text(doc):
    return "\n".join(p.text for p in doc.paragraphs)


def test_max_termin():
    rows = [{"termin": "05.10.2026"}, {"termin": "16.11.2026"}, {"termin": "01.09.2026"}]
    assert _max_termin(rows) == "16.11.2026"


def test_enddatum_wird_gefuellt():
    gen = _gen()
    answers = {"termine": {"extracted": [
        {"ergebnis": "Stakeholder-Liste", "termin": "28.09.2026"},
        {"ergebnis": "Meilenstein Durchfuehrungsfreigabe", "termin": "16.11.2026"},
    ]}}
    doc = Document(gen.generate("hermes_pia", answers, {"projektname": "T", "version": "0.1"}))
    full = _full_text(doc)
    assert "16.11.2026" in full
    assert "tt.mm.jjjj" not in full.split("Enddatum")[-1][:60]


def test_freitext_als_echte_absaetze_mit_fettem_komplexitaetskopf():
    gen = _gen()
    text = ("Basis-Ausgangslage.\n\nKomplexitätseinschätzung der Initialisierung:\n"
            "Technologie – hoch: Neuartig.\nPolitik & Stakeholder – mittel: Mehrere Ämter.")
    answers = {"ausgangslage": {"extracted": {"text": text}}}
    doc = Document(gen.generate("hermes_pia", answers, {"projektname": "T", "version": "0.1"}))
    full = _full_text(doc)
    assert "Komplexitätseinschätzung der Initialisierung" in full
    assert "Technologie" in full and "Politik & Stakeholder" in full
    # Echte Absätze statt [Shift]+[Enter]: der Kopf und die Dimensionen stehen in
    # SEPARATEN Absätzen (nicht in einem einzigen mit <w:br/>).
    W_P, W_T, W_B = f"{{{W}}}p", f"{{{W}}}t", f"{{{W}}}b"
    kopf_bold = None
    tech_eigener_absatz = False
    for p in doc.element.body.iter(W_P):
        ptxt = "".join(t.text or "" for t in p.iter(W_T))
        if "Komplexitätseinschätzung der Initialisierung" in ptxt:
            kopf_bold = any(True for _ in p.iter(W_B))
            assert "Technologie" not in ptxt        # eigener Absatz, kein Block mehr
        if ptxt.strip().startswith("Technologie"):
            tech_eigener_absatz = True
    assert kopf_bold, "Komplexitätskopf muss fett sein"
    assert tech_eigener_absatz, "Dimensionen müssen eigene Absätze sein"


def test_projektauftrag_wird_zu_durchfuehrungsauftrag():
    gen = _gen()
    answers = {"risiken": {"extracted": [
        {"beschreibung": "Risiko", "ew": "Mittel", "ag": "Hoch",
         "massnahmen": "Im Projektauftrag verankern", "verantwortung": "Projektleiter", "termin": "laufend"},
    ]}}
    doc = Document(gen.generate("hermes_pia", answers, {"projektname": "T", "version": "0.1"}))
    full_cells = " ".join(c.text for t in doc.tables for row in t.rows for c in row.cells)
    assert "Durchführungsauftrag verankern" in full_cells
    assert "Projektauftrag" not in full_cells


def test_projektorganisation_bestaetigung_in_letzter_spalte():
    gen = _gen()
    row = {"rolle_person": "Projektleiter", "bestaetigung": "ausstehend"}
    for i in range(1, 10):
        row[f"monat_{i}"] = "5" if i <= 3 else ""
    answers = {"projektorganisation": {"extracted": [row]}}
    doc = Document(gen.generate("hermes_pia", answers, {"projektname": "T", "version": "0.1"}))
    W_TR = f"{{{W}}}tr"
    for tbl in doc.element.body.iter(f"{{{W}}}tbl"):
        if "Bestätigung" in "".join(tbl.itertext()) and "Monat 1" in "".join(tbl.itertext()):
            # Datenzeile finden: letzte Zelle muss 'ausstehend' sein, Monat 1 = '5'
            from docx.table import Table
            t = Table(tbl, doc)
            data = [r for r in t.rows if "Projektleiter" in r.cells[0].text]
            assert data, "Datenzeile nicht gefunden"
            cells = [c.text.strip() for c in data[0].cells]
            assert cells[-1] == "ausstehend", cells
            assert cells[1] == "5"  # Monat 1
            break
    else:
        raise AssertionError("Projektorganisation-Tabelle nicht gefunden")


def test_deckblatt_label_weiblich():
    gen = _gen()
    md = {"projektname": "T", "version": "0.1",
          "projektleiter": "Helene Digital", "auftraggeber": "Monika Musterfrau",
          "projektleiter_geschlecht": "w", "auftraggeber_geschlecht": "w",
          "autor_geschlecht": "w", "autor": "Helene Digital"}
    doc = Document(gen.generate("hermes_pia", {}, md))
    full = _full_text(doc)
    assert "Projektleiterin" in full and "Auftraggeberin" in full and "Autorin" in full
    # Doppelform verschwindet, wenn Geschlecht bekannt
    assert "Projektleiter/in" not in full


def test_deckblatt_label_maennlich():
    gen = _gen()
    md = {"projektname": "T", "version": "0.1",
          "projektleiter": "Hans Muster", "auftraggeber": "Peter Beispiel",
          "projektleiter_geschlecht": "m", "auftraggeber_geschlecht": "m",
          "autor_geschlecht": "m", "autor": "Hans Muster"}
    doc = Document(gen.generate("hermes_pia", {}, md))
    full = _full_text(doc)
    assert "Auftraggeber" in full and "Auftraggeberin" not in full


def test_deckblatt_label_unbekannt_behaelt_doppelform():
    gen = _gen()
    md = {"projektname": "T", "version": "0.1",
          "projektleiter": "K. Broennimann", "projektleiter_geschlecht": "u",
          "autor": "K. Broennimann", "autor_geschlecht": "u"}
    doc = Document(gen.generate("hermes_pia", {}, md))
    full = _full_text(doc)
    assert "Projektleiter/in" in full or "Autor/-in" in full


def test_referenzierte_dokumente_nummer_als_textmarke():
    gen = _gen()
    answers = {"referenzierte_dokumente": {"extracted": [
        {"name": "DSG", "link": ""}, {"name": "StPO", "link": ""},
    ]}}
    doc = Document(gen.generate("hermes_pia", answers, {"projektname": "T", "version": "0.1"}))
    cells = " ".join(c.text for t in doc.tables for row in t.rows for c in row.cells)
    assert "[01]" in cells and "[02]" in cells
    # Textmarke vorhanden und referenzierbar
    names = [b.get(f"{{{W}}}name")
             for b in doc.element.body.iter(f"{{{W}}}bookmarkStart")]
    assert "ref_01" in names and "ref_02" in names

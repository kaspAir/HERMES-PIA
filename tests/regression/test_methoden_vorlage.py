"""Kundeneigene Word-Vorlage: Org-/Projekt-Ablage (Projekt übersteuert Org),
Upload-Routen und Erkennungs-Vorschau. Spiegelt die Präsentationsvorlage-Mechanik.
"""
import base64
from io import BytesIO

import pytest
from docx import Document

from app.config import Config
from app.factory import create_app


@pytest.fixture
def app(tmp_path):
    from app.shared.database import SessionLocal

    db_path = str(tmp_path / "methvorlage.db").replace("\\", "/")

    class _Cfg(Config):
        DATABASE_URL = "sqlite:///" + db_path
        SUPERADMIN_EMAIL = "betreiber@test.ch"
        SUPERADMIN_PASSWORD = "pw-super"
        SECRET_KEY = "test-secret"

    SessionLocal.remove()
    application = create_app(_Cfg)
    SessionLocal.remove()
    yield application
    SessionLocal.remove()


def _docx(headings):
    doc = Document()
    for level, text in headings:
        doc.add_heading(text, level=level)
        doc.add_paragraph("Text.")
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _b64(data):
    return base64.b64encode(data).decode("ascii")


@pytest.fixture
def ctx(app):
    """Org, Org-Admin, Nur-Leser und ein Projekt der Org."""
    auth = app.auth_service
    org = auth.create_org("Org A")
    auth.create_user("admin@a.ch", "pw", role="org_admin", org_id=org.id,
                     can_read=True, can_write=True, can_delete=True)
    auth.create_user("leser@a.ch", "pw", org_id=org.id, can_read=True)
    projekt = app.projekt_service.create_projekt(org_id=org.id, name="Projekt X")

    def _client(email):
        c = app.test_client()
        c.post("/login", data={"email": email, "password": "pw"})
        return c

    return {"app": app, "org_id": org.id, "projekt_id": projekt.id,
            "admin": _client("admin@a.ch"), "leser": _client("leser@a.ch")}


# ---- Ablage & Auflösung ------------------------------------------------ #

def test_resolve_projekt_uebersteuert_org(ctx):
    svc = ctx["app"].projekt_service
    svc.add_methoden_vorlage("org.docx", _docx([(1, "Ausgangslage")]),
                             org_id=ctx["org_id"], projekt_id=None)
    projekt = svc.get_projekt(ctx["projekt_id"])
    # Nur Org-Vorlage vorhanden -> sie greift
    assert svc.resolve_methoden_vorlage(projekt).filename == "org.docx"
    # Projekt-Vorlage übersteuert
    svc.add_methoden_vorlage("projekt.docx", _docx([(1, "Ziele")]),
                             org_id=ctx["org_id"], projekt_id=ctx["projekt_id"])
    assert svc.resolve_methoden_vorlage(projekt).filename == "projekt.docx"


def test_neuester_upload_je_scope_zaehlt(ctx):
    svc = ctx["app"].projekt_service
    svc.add_methoden_vorlage("alt.docx", _docx([(1, "Ausgangslage")]),
                             org_id=ctx["org_id"], projekt_id=None)
    svc.add_methoden_vorlage("neu.docx", _docx([(1, "Ausgangslage")]),
                             org_id=ctx["org_id"], projekt_id=None)
    assert svc.org_methoden_vorlage(ctx["org_id"]).filename == "neu.docx"


# ---- Upload-Routen ----------------------------------------------------- #

def test_pmo_upload_und_erkennungsvorschau(ctx):
    data = _docx([(1, "Ausgangssituation"), (1, "Zielsetzung"),
                  (1, "Datenschutzkonzept")])
    r = ctx["admin"].post("/pmo/methoden-vorlage",
                          json={"filename": "kunde.docx", "data": _b64(data)})
    assert r.status_code == 200 and r.get_json()["ok"] is True

    html = ctx["admin"].get("/pmo").get_data(as_text=True)
    assert "kunde.docx" in html
    # Vorschau: 2 erkannt (Ausgangslage, Ziele), 1 generisch (Datenschutzkonzept)
    assert ">2</strong> von 3" in html
    assert "als HERMES-Kapitel erkannt" in html
    assert "Datenschutzkonzept" in html


def test_projekt_upload_uebersteuert_und_zeigt_badge(ctx):
    pid = ctx["projekt_id"]
    data = _docx([(1, "Ausgangslage"), (1, "Ziele")])
    r = ctx["admin"].post(f"/projekt/{pid}/methoden-vorlage",
                          json={"filename": "proj.docx", "data": _b64(data)})
    assert r.status_code == 200
    html = ctx["admin"].get(f"/projekt/{pid}").get_data(as_text=True)
    assert "proj.docx" in html


def test_upload_validiert_dateityp(ctx):
    r = ctx["admin"].post("/pmo/methoden-vorlage",
                          json={"filename": "x.pptx", "data": _b64(b"PK\x03\x04stuff")})
    assert r.status_code == 400
    # .dotx ist erlaubt (Template-Datei)
    r2 = ctx["admin"].post("/pmo/methoden-vorlage",
                           json={"filename": "x.dotx", "data": _b64(_docx([(1, "Ziele")]))})
    assert r2.status_code == 200


def test_leser_darf_nicht_hochladen(ctx):
    pid = ctx["projekt_id"]
    data = _b64(_docx([(1, "Ziele")]))
    assert ctx["leser"].post("/pmo/methoden-vorlage",
                             json={"filename": "x.docx", "data": data}).status_code == 403
    assert ctx["leser"].post(f"/projekt/{pid}/methoden-vorlage",
                             json={"filename": "x.docx", "data": data}).status_code == 403


def test_ohne_vorlage_kein_erkennungsblock(ctx):
    html = ctx["admin"].get("/pmo").get_data(as_text=True)
    assert "als HERMES-Kapitel erkannt" not in html
    assert "das Interview folgt der HERMES-Standardstruktur" in html

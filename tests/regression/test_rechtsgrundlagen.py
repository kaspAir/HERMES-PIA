"""Beweist: Rechtsgrundlagenanalyse (Phase A) – Seeding aus PIA, beratender Entwurf,
Befüllung ins HERMES-Template. Modular, ohne den PIA zu berühren."""
import json

import pytest
from docx import Document

from app.config import Config
from app.domains.ergebnisse.projektwissen import Projektwissen
from app.domains.ergebnisse.rechtsgrundlagen.service import RechtsgrundlagenService
from app.factory import create_app


_PIA = {
    "ausgangslage": {"extracted": {"text": "Ablösung des Justizsystems Juris Fiat."}},
    "referenzierte_dokumente": {"extracted": [
        {"nr": "01", "name": "Schweizerische Strafprozessordnung (StPO)", "link": ""},
        {"nr": "02", "name": "Bundesgesetz über die Produktehaftpflicht (PrHG)", "link": ""},
        {"nr": "03", "name": "Bundesgesetz über den Datenschutz (DSG)", "link": ""},  # -> Schuban
    ]},
    "mitgeltende_unterlagen": {"extracted": [
        {"name": "Kantonales Beschaffungsrecht (Submissionsgesetz)", "link": ""}]},  # kantonal
    "ziele": {"extracted": [{"beschreibung": "Verfahrensschritte im System abbilden"}]},
}


class _FakeLLM:
    def __init__(self, payload):
        self._payload = payload

    def complete(self, system, messages, max_tokens=1024):
        return json.dumps(self._payload)


class _FakeFedlex:
    """Kein Netzwerk: liefert die vorgegebene {Suchbegriff: [Treffer]}-Zuordnung."""
    def __init__(self, mapping=None):
        self._m = mapping or {}

    def suche_mehrere(self, begriffe, **kw):
        return self._m


# ---- Modul-Logik (kein DB, kein echter LLM) ------------------------------- #

def test_seeding_ohne_llm_uebernimmt_pia_gesetze():
    wissen = Projektwissen(_PIA, ebene="kanton")   # ohne konkreten Kanton -> kein Sammlungslink
    svc = RechtsgrundlagenService(None, None, None, llm=None, fedlex=_FakeFedlex())
    answers = svc.build_answers(wissen)
    namen = [r["rechtsgrundlage"] for r in answers["bestehende_rechtsgrundlagen"]["extracted"]]
    assert "Schweizerische Strafprozessordnung (StPO)" in namen
    # Datenschutzgesetz gehört in die Schutzbedarfsanalyse -> hier NICHT
    assert "Bundesgesetz über den Datenschutz (DSG)" not in namen
    # Ohne LLM keine Beschreibung erfunden
    assert all(r.get("beschreibung", "") == ""
               for r in answers["bestehende_rechtsgrundlagen"]["extracted"])
    # 0.4 Definitionen: enthält die Kürzel der genannten Gesetze
    abk = [r["abkuerzung"] for r in answers["definitionen"]["extracted"]]
    assert "StPO" in abk and "PrHG" in abk
    # Frueher: leerer Text. Ein leeres Pflichtkapitel liest sich aber wie
    # «geprueft und unbedenklich» - jetzt sagt es, dass es fehlt.
    assert "nicht beurteilt" in answers["konsequenzen"]["extracted"]["text"]


def test_datenschutz_und_ebene_filter():
    svc = RechtsgrundlagenService(None, None, None, llm=None, fedlex=_FakeFedlex())
    # Nur Bund: kantonale Gesetze fallen weg; Datenschutz immer weg
    namen_bund = [r["rechtsgrundlage"] for r in
                  svc.build_answers(Projektwissen(_PIA, ebene="bund"))
                  ["bestehende_rechtsgrundlagen"]["extracted"]]
    assert "Schweizerische Strafprozessordnung (StPO)" in namen_bund
    assert "Kantonales Beschaffungsrecht (Submissionsgesetz)" not in namen_bund   # kantonal, nur Bund
    assert "Bundesgesetz über den Datenschutz (DSG)" not in namen_bund            # Datenschutz -> Schuban
    # Bund + Kanton: kantonales Gesetz wieder dabei
    namen_beide = [r["rechtsgrundlage"] for r in
                   svc.build_answers(Projektwissen(_PIA, ebene="bund,kanton"))
                   ["bestehende_rechtsgrundlagen"]["extracted"]]
    assert "Kantonales Beschaffungsrecht (Submissionsgesetz)" in namen_beide


def test_llm_vorschlag_wird_gemergt_und_pia_bleibt_fuehrend():
    wissen = Projektwissen(_PIA, ebene="bund")
    svc = RechtsgrundlagenService(None, None, None, fedlex=_FakeFedlex(), llm=_FakeLLM({
        "bestehende": [{"rechtsgrundlage": "Schweizerische Strafprozessordnung (StPO)",
                        "beschreibung": "Regelt das Strafverfahren."}],
        "luecken": [{"luecke": "Fehlende Grundlage", "beschreibung": "für neue Bearbeitung"}],
        "konsequenzen": "Rechtliches Risiko ohne Anpassung.",
        "empfehlung": "Rechtsgrundlage vor Realisierung schaffen.",
    }))
    answers = svc.build_answers(wissen)
    best = {r["rechtsgrundlage"]: r.get("beschreibung", "")
            for r in answers["bestehende_rechtsgrundlagen"]["extracted"]}
    assert best["Schweizerische Strafprozessordnung (StPO)"] == "Regelt das Strafverfahren."
    assert "Bundesgesetz über die Produktehaftpflicht (PrHG)" in best   # weiteres Bundesgesetz
    luecken = answers["identifizierte_luecken"]["extracted"]
    assert luecken and luecken[0]["luecke"] == "Fehlende Grundlage"
    assert answers["konsequenzen"]["extracted"]["text"].startswith("Rechtliches Risiko")


# ---- End-to-End: Entwurf aus PIA-Session -> .docx ------------------------- #

@pytest.fixture
def app(tmp_path):
    from app.shared.database import SessionLocal
    db_path = str(tmp_path / "rga.db").replace("\\", "/")

    class _Cfg(Config):
        DATABASE_URL = "sqlite:///" + db_path
        SECRET_KEY = "x"

    SessionLocal.remove()
    application = create_app(_Cfg)
    SessionLocal.remove()
    yield application
    SessionLocal.remove()


def _projekt_mit_pia(app):
    from app.domains.interview.models import InterviewSession
    from app.shared.database import SessionLocal
    ps = app.projekt_service
    projekt = ps.create_projekt(org_id=1, name="BKI Test 2", auftraggeber="Monika Musterfrau")
    erg = ps.add_ergebnis(projekt.id, "projektinitialisierungsauftrag", created_by="Helene Digital")
    db = SessionLocal()
    db.add(InterviewSession(
        method_id="hermes_pia", project_name="BKI Test 2", org_id=1,
        created_by="Helene Digital", auftraggeber="Monika Musterfrau",
        ergebnis_id=erg.id, answers_json=json.dumps(_PIA)))
    db.commit()
    return projekt.id


def test_entwurf_und_docx_end_to_end(app):
    with app.app_context():
        pid = _projekt_mit_pia(app)
        svc = app.rechtsgrundlagen_service
        projekt = app.projekt_service.get_projekt(pid)

        entwurf = svc.erzeuge_entwurf(projekt, ebene="kanton", kanton="ZH")
        assert entwurf.answers_json and entwurf.kanton == "ZH"

        doc = Document(svc.generate_docx(projekt))
        cells = " ".join(c.text for t in doc.tables for row in t.rows for c in row.cells)
        # Aus dem PIA übernommenes (Nicht-Datenschutz-)Gesetz steht im Dokument
        assert "Schweizerische Strafprozessordnung (StPO)" in cells
        # Titel/Kapitel des Templates vorhanden
        volltext = "\n".join(p.text for p in doc.paragraphs)
        assert "Bestehende Rechtsgrundlagen" in volltext
        assert "Empfehlung" in volltext


# ---- Phase B: Fedlex-Grounding (ohne Netzwerk) ---------------------------- #

def test_fedlex_offline_index_lookup():
    from app.domains.rechtsquellen.fedlex import FedlexClient
    # Injizierter Mini-Index (kein Netzwerk): kürzeste SR = Haupterlass, Wortgrenzen.
    client = FedlexClient(index=[
        {"sr": "312.0", "titel": "Schweizerische Strafprozessordnung vom 5. Oktober 2007 (StPO)",
         "url": "https://www.fedlex.admin.ch/eli/cc/2010/267/de"},
        {"sr": "312.1", "titel": "Schweizerische Jugendstrafprozessordnung (JStPO)",
         "url": "https://www.fedlex.admin.ch/eli/cc/2010/226/de"},
        {"sr": "272", "titel": "Bundesgesetz über den Gerichtsstand (Gerichtsstandsgesetz)",
         "url": "x"},
    ])
    res = client.suche_mehrere(["Strafprozessordnung", "DSG"])
    assert res["Strafprozessordnung"][0]["sr"] == "312.0"
    assert res["Strafprozessordnung"][0]["url"].endswith("/de")
    assert "DSG" not in res             # 'dsg' matcht NICHT in 'GerichtsstanDSGesetz'


def test_fedlex_index_mitgeliefert():
    """Der reale Offline-Index ist vorhanden und findet Kernerlasse."""
    from app.domains.rechtsquellen.fedlex import FedlexClient
    res = FedlexClient().suche_mehrere(["Strafprozessordnung", "Datenschutz"])
    assert res["Strafprozessordnung"][0]["sr"] == "312.0"
    assert res["Datenschutz"][0]["sr"] == "235.1"
    assert "fedlex.admin.ch" in res["Strafprozessordnung"][0]["url"]


def test_suchbegriffe_aus_gesetzesname():
    from app.domains.ergebnisse.rechtsgrundlagen.grounding import suchbegriffe
    t = suchbegriffe("Bundesgesetz über den Datenschutz (DSG)")
    assert "DSG" in t and "Datenschutz" in t
    assert "Bundesgesetz" not in t                    # generisch -> nicht als Begriff


def test_ground_federal_immer(_ebene="egal"):
    from app.domains.ergebnisse.rechtsgrundlagen.grounding import ground_federal
    fake = _FakeFedlex({"StPO": [{"sr": "312.0", "titel": "StPO …", "url": "u"}]})
    namen = ["Schweizerische Strafprozessordnung (StPO)"]
    # Bundesrecht gilt in jedem Kanton -> auch bei Kantonsebene grounden.
    for ebene in ("kanton", "bund", "bund,kanton"):
        g = ground_federal(namen, ebene, fake)
        assert g["Schweizerische Strafprozessordnung (StPO)"]["sr"] == "312.0"


def test_service_reichert_referenzierte_und_bestehende_mit_fundstelle_an():
    fake = _FakeFedlex({
        "StPO": [{"sr": "312.0", "titel": "Schweizerische Strafprozessordnung vom 5. Oktober 2007 (StPO)",
                  "url": "https://www.fedlex.admin.ch/eli/cc/2010/267/de"}],
    })
    wissen = Projektwissen(_PIA, ebene="bund")
    svc = RechtsgrundlagenService(None, None, None, llm=None, fedlex=fake)
    answers = svc.build_answers(wissen)
    # Referenzierte: der StPO-Eintrag bekommt SR + Fedlex-Link
    ref = next(r for r in answers["referenzierte_dokumente"]["extracted"]
               if "Strafprozessordnung" in r["name"])
    assert "SR 312.0" in ref["link"] and "fedlex.admin.ch" in ref["link"]
    # Bestehende Rechtsgrundlagen: verifizierter Titel + SR + Link als Beschreibung
    best = next(r for r in answers["bestehende_rechtsgrundlagen"]["extracted"]
                if "Strafprozessordnung" in r["rechtsgrundlage"])
    assert "SR 312.0" in best["beschreibung"] and "fedlex.admin.ch" in best["beschreibung"]


def test_llm_entdeckt_zusaetzliches_gesetz_und_weist_keine_luecke_aus():
    wissen = Projektwissen(_PIA, ebene="bund")
    svc = RechtsgrundlagenService(None, None, None, fedlex=_FakeFedlex(), llm=_FakeLLM({
        "bestehende": [{"rechtsgrundlage": "Strafregistergesetz (StReG)",
                        "beschreibung": "Regelt das Strafregister VOSTRA."}],
        "luecken": [],   # keine Lücke
    }))
    answers = svc.build_answers(wissen)
    namen = [r["rechtsgrundlage"] for r in answers["bestehende_rechtsgrundlagen"]["extracted"]]
    assert "Strafregistergesetz (StReG)" in namen          # vom LLM selbst gefunden (nicht im PIA)
    luecken = answers["identifizierte_luecken"]["extracted"]
    # Frueher: «Keine Lücke identifiziert» samt Behauptung, es bestehe eine
    # Rechtsgrundlage. Diese Analyse prueft die Ziele aber nicht einzeln - sie
    # darf deshalb keine Entwarnung geben.
    assert luecken and luecken[0]["luecke"] == "Nicht abschliessend geprüft"


def test_betriebskonzept_ist_keine_rechtsgrundlage():
    pia = {"referenzierte_dokumente": {"extracted": [
        {"name": "Betriebskonzept / Systemdokumentation Juris Fiat", "link": ""},
        {"name": "Schweizerische Strafprozessordnung (StPO)", "link": ""}]}}
    svc = RechtsgrundlagenService(None, None, None, llm=None, fedlex=_FakeFedlex())
    namen = [r["rechtsgrundlage"] for r in
             svc.build_answers(Projektwissen(pia, ebene="bund"))
             ["bestehende_rechtsgrundlagen"]["extracted"]]
    assert "Schweizerische Strafprozessordnung (StPO)" in namen
    assert not any("Betriebskonzept" in n for n in namen)   # kein Gesetz -> raus aus Kap.1


def test_product_compliance_ohne_datenschutz_infosec():
    wissen = Projektwissen(_PIA, ebene="bund")
    svc = RechtsgrundlagenService(None, None, None, fedlex=_FakeFedlex(), llm=_FakeLLM({
        "compliance": [
            {"compliance": "Ausschreibungspflicht (BöB/VöB)", "beschreibung": "Beschaffung ..."},
            {"compliance": "Datenschutz (DSG/VDSG)", "beschreibung": "besonders schützenswerte Daten"},
            {"compliance": "Informationssicherheit (ISG/ISV)", "beschreibung": "klassifizierte Daten"},
        ],
    }))
    comp = [r["compliance"] for r in svc.build_answers(wissen)["product_compliance"]["extracted"]]
    assert any("Ausschreibungspflicht" in c for c in comp)
    # Datenschutz/Infosec gehören in die Schutzbedarfsanalyse -> hier raus
    assert not any("Datenschutz" in c or "Informationssicherheit" in c for c in comp)


def test_hermes_guardrails_im_prompt():
    from app.domains.ergebnisse.rechtsgrundlagen.proposals import SYSTEM
    s = SYSTEM.lower()
    assert "schutzbedarfsanalyse" in s and "isds" in s
    assert "weiteres vorgehen" in s and "klassisch/agil" in s


def test_kantonaler_sammlungslink_fuer_ungegroundete_gesetze():
    from app.domains.rechtsquellen.kantone import sammlung_link
    assert "zh.ch" in sammlung_link("ZH").lower()
    pia = {"referenzierte_dokumente": {"extracted": [
        {"name": "Kantonales Justizvollzugsgesetz", "link": ""},
        {"name": "Schweizerische Strafprozessordnung (StPO)", "link": ""}]}}
    fake = _FakeFedlex({"StPO": [{"sr": "312.0",
        "titel": "Schweizerische Strafprozessordnung (StPO)",
        "url": "https://www.fedlex.admin.ch/eli/cc/2010/267/de"}]})
    svc = RechtsgrundlagenService(None, None, None, llm=None, fedlex=fake)
    answers = svc.build_answers(Projektwissen(pia, ebene="bund,kanton", kanton="ZH"))
    ref = {r["name"]: r["link"] for r in answers["referenzierte_dokumente"]["extracted"]}
    # Bundesgesetz -> Fedlex-SR; kantonales Gesetz -> Link zur ZH-Sammlung
    assert "SR 312.0" in ref["Schweizerische Strafprozessordnung (StPO)"]
    assert "Kantonale Sammlung ZH" in ref["Kantonales Justizvollzugsgesetz"]


def test_kein_kantonslink_ohne_kantonsebene():
    pia = {"referenzierte_dokumente": {"extracted": [
        {"name": "Kantonales Polizeigesetz", "link": ""}]}}
    svc = RechtsgrundlagenService(None, None, None, llm=None, fedlex=_FakeFedlex())
    answers = svc.build_answers(Projektwissen(pia, ebene="bund", kanton="ZH"))  # nur Bund
    # Kein Kantonslink, wenn Kantonsebene nicht gewählt (kantonales Gesetz zudem gefiltert)
    ref = answers["referenzierte_dokumente"]["extracted"]
    assert all("Kantonale Sammlung" not in r["link"] for r in ref)


# ---- Leere Kapitel sagen, WARUM sie leer sind ---------------------------- #

def _svc():
    from app.domains.ergebnisse.rechtsgrundlagen.service import RechtsgrundlagenService
    return RechtsgrundlagenService.__new__(RechtsgrundlagenService)


def test_leeres_kapitel_liefert_nie_eine_leere_zeile():
    """Gemessen an einer echten Analyse (Testprojekt 17): «Bevorstehende
    Änderungen» und «Vorschläge zur Deckung» enthielten eine Zeile mit der
    Nummer «01» und sonst nichts. Die Vorlage nummeriert die Zeile – für den
    Leser sieht das aus, als sei die Erzeugung abgebrochen."""
    svc = _svc()
    zeilen = svc._rows_or_blank(None, ("rechtsgrundlage", "beschreibung", "auswirkung"),
                                "bevorstehende_aenderungen")
    assert len(zeilen) == 1
    assert zeilen[0]["rechtsgrundlage"].startswith("Keine bevorstehende")
    assert zeilen[0]["beschreibung"]
    # Geprüft-und-nichts-gefunden darf nicht wie Entwarnung klingen.
    assert "nicht ausgeschlossen" in zeilen[0]["beschreibung"]


def test_product_compliance_leer_wird_benannt():
    svc = _svc()
    z = svc._rows_or_blank([], ("compliance", "beschreibung"), "product_compliance")
    assert z[0]["compliance"] == "Kein Hinweis identifiziert"
    assert z[0]["beschreibung"]


def test_ohne_luecke_entfaellt_die_deckung():
    """Ohne Lücke gibt es nichts zu decken – das ist ein Ergebnis, kein Ausfall."""
    svc = _svc()
    luecken = [{"luecke": "Keine Lücke identifiziert", "beschreibung": "…"}]
    z = svc._deckungsvorschlaege(None, luecken)
    assert z[0]["luecke"] == "Entfällt"
    assert "keine Lücken" in z[0]["vorschlag"]


def test_luecke_ohne_vorschlag_bleibt_sichtbar_offen():
    """Gibt es eine Lücke, aber keinen Vorschlag, ist die Frage OFFEN – und muss
    so dastehen. Eine leere Zeile liesse «geprüft» und «unbeantwortet» gleich
    aussehen."""
    svc = _svc()
    luecken = [{"luecke": "Keine Grundlage für die Bekanntgabe an Dritte",
                "beschreibung": "…"}]
    z = svc._deckungsvorschlaege([], luecken)
    assert len(z) == 1
    assert z[0]["luecke"].startswith("Keine Grundlage")
    assert z[0]["vorschlag"].startswith("Offen")


def test_vorhandene_vorschlaege_bleiben_unangetastet():
    svc = _svc()
    echte = [{"luecke": "L", "vorschlag": "Verordnung anpassen"}]
    assert svc._deckungsvorschlaege(echte, []) == echte


# ---- Keine unverdiente Entwarnung ---------------------------------------- #

def test_verfassung_und_emrk_sind_keine_ermaechtigungsgrundlage():
    """Gemessen (BKI Test 6, biometrische Massenüberwachung): BV und EMRK
    standen als «Bestehende Rechtsgrundlage» Nr. 01 und 02. Das ist eine
    Umkehrung – sie sind die SCHRANKE des Eingriffs, nicht seine Ermächtigung.
    Ein Dokument, das sie so aufführt, liest sich wie eine Erlaubnis."""
    svc = _svc()
    for name in ("Bundesverfassung der Schweizerischen Eidgenossenschaft (BV)",
                 "Europäische Menschenrechtskonvention (EMRK)",
                 "Kantonsverfassung des Kantons St. Gallen"):
        assert svc._ist_schrankennorm(name), name
        assert not svc._kap1_geeignet(name, "Bund und Kanton")
    # Ein echtes Ermächtigungsgesetz bleibt unberührt.
    assert not svc._ist_schrankennorm("Polizeigesetz des Kantons St. Gallen")


def test_schrankennorm_als_grundlage_wird_zur_luecke():
    """Sie darf nicht still verschwinden: dass nur Schrankennormen genannt
    wurden, IST der Befund – es wurde keine Ermächtigung gefunden."""
    svc = _svc()
    rows = svc._luecken(None, ["Europäische Menschenrechtskonvention (EMRK)"])
    assert len(rows) == 1
    assert "Keine Ermächtigungsgrundlage" in rows[0]["luecke"]
    assert "Art. 36 Abs. 1 BV" in rows[0]["beschreibung"]


def test_ohne_befund_keine_behauptung_einer_rechtsgrundlage():
    """Der alte Satz «Für die im Projekt geplanten Tätigkeiten besteht nach
    dieser Analyse eine Rechtsgrundlage» war eine nie geprüfte Behauptung –
    und stand über einem Vorhaben zur Massenüberwachung."""
    svc = _svc()
    rows = svc._luecken(None)
    assert rows[0]["luecke"] == "Nicht abschliessend geprüft"
    text = rows[0]["beschreibung"]
    assert "besteht" not in text.split("darf")[0] or "nicht" in text
    assert "nicht geschlossen werden" in text
    assert "je Projektziel" in text


def test_pflichtkapitel_bleiben_nie_wortlos():
    """Im gemessenen Lauf waren «Beurteilung der Konsequenzen» und
    «Empfehlung» vollständig leer – das Dokument sah abgeschlossen aus."""
    svc = _svc()
    assert "nicht freigabefähig" in svc._pflichttext("", "x ist nicht freigabefähig")
    assert svc._pflichttext("Echte Beurteilung.", "ersatz") == "Echte Beurteilung."

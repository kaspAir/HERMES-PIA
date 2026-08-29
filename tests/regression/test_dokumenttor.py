"""Beweist: die Beispiel-Formatvorlage bleibt nicht stehen, und das
Qualitätstor steht vor JEDEM Dokument.

Beide Befunde stammen aus einem echten Lauf und hängen zusammen. Der Generator
schrieb Inhalt in die Zelle, liess aber die Formatvorlage `HTabBeispiel85ptF`
darauf stehen — die Beispiel-Vorlage. Die Prüfung liest für D-004 die
FORMATVORLAGE, nicht den Inhalt, und meldete deshalb «Das Kapitel wurde nicht
befüllt» über «21.09.2026» und «MS Office Palette gem. Vorlage». Beide echt,
beide als Muss gemeldet, beide nicht behebbar.

Weil D-004 ein Muss ist, blockierte das den Download des Auftrags — und man
musste jedes Mal «Trotzdem herunterladen» klicken. Ein Tor, das man bei jedem
Dokument umgeht, schützt nichts mehr. Umgekehrt gingen Rechtsgrundlagenanalyse,
Checkliste und Liste Projektentscheide ganz ohne Prüfung hinaus.
"""
import io

import pytest
from docx import Document
from lxml import etree

from app.config import Config
from app.domains.generation.service import (
    _clear_info_style, _set_p_multiline, _set_p_text)
from app.domains.qualitaet.dokument import pruefe_dokument
from app.factory import create_app

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _absatz(stil):
    p = etree.Element(f"{{{W}}}p")
    pPr = etree.SubElement(p, f"{{{W}}}pPr")
    ps = etree.SubElement(pPr, f"{{{W}}}pStyle")
    ps.set(f"{{{W}}}val", stil)
    return p


def _stil(p):
    pPr = p.find(f"{{{W}}}pPr")
    if pPr is None:
        return None
    ps = pPr.find(f"{{{W}}}pStyle")
    return ps.get(f"{{{W}}}val") if ps is not None else None


# ---- Die Formatvorlage ---------------------------------------------------- #

def test_beim_schreiben_faellt_die_beispielvorlage_weg():
    """Die echte Formatvorlage aus der HERMES-Vorlage."""
    p = _absatz("HTabBeispiel85ptF")
    _set_p_text(p, "21.09.2026")
    assert _stil(p) is None
    assert "".join(t.text for t in p.iter(f"{{{W}}}t")) == "21.09.2026"


def test_eine_echte_formatvorlage_bleibt():
    """Geräumt wird NUR, was Hilfe/Beispiel/Platzhalter markiert."""
    p = _absatz("HTabText85ptF")
    _set_p_text(p, "Inhalt")
    assert _stil(p) == "HTabText85ptF"


def test_auch_die_folgeabsaetze_erben_sie_nicht():
    """Mehrzeiliger Text erzeugt neue Absätze aus demselben pPr — ohne
    Räumung trüge jede Folgezeile die Beispiel-Vorlage weiter."""
    eltern = etree.Element(f"{{{W}}}body")
    p = _absatz("HTabBeispiel85ptF")
    eltern.append(p)
    _set_p_multiline(p, "erste Zeile\nzweite Zeile\ndritte Zeile")
    absaetze = list(eltern.iter(f"{{{W}}}p"))
    assert len(absaetze) == 3
    assert all(_stil(a) is None for a in absaetze)


@pytest.mark.parametrize("stil", [
    "HTabBeispiel85ptF", "Hilfetext", "InfoText", "PlatzhalterZeile", "MusterZeile"])
def test_alle_hilfsvorlagen_werden_erkannt(stil):
    p = _absatz(stil)
    _clear_info_style(p)
    assert _stil(p) is None


# ---- Das Tor vor jedem Dokument ------------------------------------------- #

@pytest.fixture
def app(tmp_path):
    from app.shared.database import SessionLocal

    class _Cfg(Config):
        DATABASE_URL = "sqlite:///" + str(tmp_path / "tor.db").replace("\\", "/")
        SECRET_KEY = "test-secret"

    SessionLocal.remove()
    anwendung = create_app(_Cfg)
    SessionLocal.remove()
    yield anwendung
    SessionLocal.remove()


@pytest.fixture
def projekt(app):
    auth = app.auth_service
    org_id = auth.create_org("Org").id
    auth.create_user("pl@org.ch", "pw", org_id=org_id, can_read=True, can_write=True)
    c = app.test_client()
    c.post("/login", data={"email": "pl@org.ch", "password": "pw"})
    c.post("/interview/start", data={"project_name": "Musterprojekt",
                                     "projektleiter": "Frau Muster"})
    return c, app.projekt_service.projekte_for_org(org_id)[0]


def _mit_platzhalter():
    """Ein Dokument, das einen Datums-Platzhalter der Vorlage trägt."""
    dok = Document()
    t = dok.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "Version"
    t.cell(0, 1).text = "Datum"
    t.cell(1, 0).text = "0.1"
    t.cell(1, 1).text = "tt.mm.jjjj"
    puffer = io.BytesIO()
    dok.save(puffer)
    puffer.seek(0)
    return puffer


def test_der_platzhalter_wird_ueberhaupt_erkannt():
    """Die Gegenprobe zum Tor: ohne diesen Befund prüfte es nichts."""
    befunde = pruefe_dokument(Document(_mit_platzhalter()))
    assert any(b.regel == "D-003" for b in befunde)


def test_die_rechtsgrundlagenanalyse_wird_gesperrt(app, projekt, monkeypatch):
    c, p = projekt
    monkeypatch.setattr(app.rechtsgrundlagen_service, "generate_docx",
                        lambda projekt: _mit_platzhalter(), raising=False)
    antwort = c.get(f"/projekt/{p.id}/rechtsgrundlagen/download/x.docx")
    assert antwort.status_code == 409
    assert "D-003" in antwort.data.decode("utf-8")


def test_der_notausgang_liefert_trotzdem_aus(app, projekt, monkeypatch):
    """Er gehört dazu – aber als Ausnahme, nicht als Weg."""
    c, p = projekt
    monkeypatch.setattr(app.rechtsgrundlagen_service, "generate_docx",
                        lambda projekt: _mit_platzhalter(), raising=False)
    antwort = c.get(f"/projekt/{p.id}/rechtsgrundlagen/download/x.docx?trotzdem=1")
    assert antwort.status_code == 200
    assert antwort.data[:2] == b"PK"


def test_ein_sauberes_dokument_geht_durch(app, projekt, monkeypatch):
    c, p = projekt
    dok = Document()
    dok.add_paragraph("Alles befüllt.")
    puffer = io.BytesIO()
    dok.save(puffer)
    puffer.seek(0)
    monkeypatch.setattr(app.rechtsgrundlagen_service, "generate_docx",
                        lambda projekt: puffer, raising=False)
    antwort = c.get(f"/projekt/{p.id}/rechtsgrundlagen/download/x.docx")
    assert antwort.status_code == 200


# ---- Geltungsbereich: eine Regel, die überall feuert, prüft nichts ------- #

def test_der_phasenbericht_gilt_nur_fuer_die_initialisierung():
    """Die Liste Projektentscheide Steuerung führt das GANZE Projekt auf. In
    der Konzept- und Realisierungsphase ist ein Phasenbericht ein richtiges
    HERMES-Ergebnis. Die Regel dort anzuwenden hiess, der Vorlage einen Fehler
    vorzuwerfen, den sie nicht macht."""
    dok = Document()
    dok.add_paragraph("Entscheid Phasenfreigabe – Grundlage: Phasenbericht")
    regeln = {b.regel for b in pruefe_dokument(dok)}
    assert "D-005" in regeln
    ueber = {b.regel for b in pruefe_dokument(dok, phasenuebergreifend=True)}
    assert "D-005" not in ueber


def test_falsche_begriffe_bleiben_auch_phasenuebergreifend_falsch():
    """Ausgenommen ist NUR, was an der Phase hängt – nicht die ganze Regel."""
    dok = Document()
    dok.add_paragraph("Der Lenkungsausschuss entscheidet.")
    assert any(b.regel == "D-005"
               for b in pruefe_dokument(dok, phasenuebergreifend=True))


def test_das_datum_der_pruefung_ist_kein_platzhalter():
    """Drei gleich gebaute Tabellen untereinander: Änderungskontrolle, Prüfung,
    Freigabe. In der Zeile steht nur «0.1 | | tt.mm.jjjj» – welche es ist, sagt
    allein die Überschrift darüber. Ohne diesen Bezug wurde «tt.mm.jjjj» dort
    gemeldet, wo es hingehört: beide Tabellen werden erst später ausgefüllt."""
    dok = Document()
    dok.add_paragraph("Prüfung")
    t = dok.add_table(rows=2, cols=3)
    t.cell(0, 0).text = "Version"
    t.cell(1, 0).text = "0.1"
    t.cell(1, 2).text = "tt.mm.jjjj"
    assert not [b for b in pruefe_dokument(dok) if b.regel == "D-003"]


def test_das_datum_der_aenderungskontrolle_ist_einer():
    """Die Gegenprobe – sonst prüfte der Test darüber nichts."""
    dok = Document()
    dok.add_paragraph("Änderungskontrolle")
    t = dok.add_table(rows=2, cols=3)
    t.cell(0, 0).text = "Version"
    t.cell(1, 0).text = "0.1"
    t.cell(1, 2).text = "tt.mm.jjjj"
    assert [b for b in pruefe_dokument(dok) if b.regel == "D-003"]

"""Inkrement 4: Ausgabe im Format der hochgeladenen Kundenvorlage.

Nur Kapitel mit Antworten werden befüllt (Zuordnung über die Original-Überschrift);
Tabellen über die Spaltenüberschriften (inkl. Abkürzungs-Aliase). Fixe Hilfstexte
bleiben unangetastet.
"""
from io import BytesIO

from docx import Document

from app.config import Config
from app.domains.generation.service import GenerationService
from app.domains.method.service import MethodService
from app.domains.method.template_structure import build_derived_method
from app.shared.config_loader import load_method


def _canonical():
    return load_method(Config.METHODS_DIR, "hermes_pia")


def _template():
    """HERMES-Variante: Hilfstext + umbenannte Kapitel + Risiken-Tabelle mit
    abgekürzten Spaltenüberschriften (EW/AG/RZ/Verantw.)."""
    doc = Document()
    doc.add_heading("Hinweise zur Anwendung des Dokumentes", level=1)
    doc.add_paragraph("FIXER-HILFETEXT-UNVERAENDERBAR.")
    doc.add_heading("Ausgangssituation", level=1)          # -> ausgangslage
    doc.add_paragraph("Bitte hier die Ausgangslage eintragen.")
    doc.add_heading("Risikoanalyse", level=1)              # -> risiken
    t = doc.add_table(rows=2, cols=8)
    for c, h in zip(t.rows[0].cells,
                    ["Nr.", "Risikobeschreibung", "EW", "AG", "RZ",
                     "Massnahmen", "Verantw.", "Termin"]):
        c.text = h
    for c in t.rows[1].cells:
        c.text = ""
    doc.add_heading("Nicht erfragtes Kapitel", level=1)
    doc.add_paragraph("SOLL-BLEIBEN.")
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _text(doc):
    para = "\n".join(p.text for p in doc.paragraphs)
    tbl = " | ".join(c.text for t in doc.tables for r in t.rows for c in r.cells)
    return para + " || " + tbl, tbl


def test_generate_into_template_fuellt_und_schont_hilfstext():
    tpl = _template()
    method, _ = build_derived_method(tpl, _canonical())
    answers = {
        "ausgangslage": {"extracted": {"text": "MEINE-AUSGANGSLAGE"}},
        "risiken": {"extracted": [{
            "beschreibung": "RISIKO-X", "ew": "Hoch", "ag": "Mittel",
            "massnahmen": "MASSNAHME-X", "verantwortung": "PL",
            "termin": "01.01.2027"}]},
    }
    gen = GenerationService(MethodService(Config.METHODS_DIR))
    buf = gen.generate_into_template(
        tpl, method, answers, {"projektname": "P", "projektleiter": "K"})
    both, tbl = _text(Document(buf))

    # Freitext + Tabelle gefüllt
    assert "MEINE-AUSGANGSLAGE" in both
    assert "RISIKO-X" in both and "MASSNAHME-X" in both
    # EW/AG über Alias erkannt -> Risikozahl berechnet (Hoch×Mittel = 3×2 = 6)
    assert "6" in tbl
    # Instruktionsprosa des Freitextkapitels wurde ersetzt
    assert "Bitte hier die Ausgangslage eintragen" not in both
    # Fixer Hilfstext (Vorspann) unangetastet
    assert "FIXER-HILFETEXT-UNVERAENDERBAR" in both
    # Kapitel ohne Antwort bleibt unberührt
    assert "SOLL-BLEIBEN" in both


def test_tabelle_ohne_passende_spalten_bleibt_unangetastet():
    # Eine Tabelle, deren Überschriften zu KEINER Spalte passen, wird nicht
    # korrumpiert (sondern in Ruhe gelassen).
    from docx import Document as _D
    doc = _D()
    doc.add_heading("Ausgangssituation", level=1)
    doc.add_paragraph("x")
    doc.add_heading("Risikoanalyse", level=1)
    t = doc.add_table(rows=2, cols=2)
    for c, h in zip(t.rows[0].cells, ["Voellig", "Fremd"]):
        c.text = h
    t.rows[1].cells[0].text = "URSPRUNG-A"
    t.rows[1].cells[1].text = "URSPRUNG-B"
    buf = BytesIO(); doc.save(buf); tpl = buf.getvalue()

    method, _ = build_derived_method(tpl, _canonical())
    answers = {"risiken": {"extracted": [{"beschreibung": "R"}]}}
    gen = GenerationService(MethodService(Config.METHODS_DIR))
    out = Document(gen.generate_into_template(tpl, method, answers, {}))
    both, _ = _text(out)
    assert "URSPRUNG-A" in both and "URSPRUNG-B" in both
